"""
Professional Project Report — Scholarly Library Management System
Screenshots are embedded inline within each relevant section.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SHOT_DIR    = r"d:\Learn\Library management app\screenshots"
OUTPUT_PATH = r"d:\Learn\Library management app\Project_Report_Scholarly_Library.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)

# ── Low-level helpers ─────────────────────────────────────────────────────────
def sf(run, size, bold=False, italic=False, name="Times New Roman"):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    try: run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    except: pass

def _new_para(align, sb, sa):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def para(text, size=12, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = _new_para(align, sb, sa)
    if text:
        r = p.add_run(text)
        sf(r, size, bold, italic)
    return p

def mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = _new_para(align, sb, sa)
    for text, bold, italic, size in parts:
        r = p.add_run(text)
        sf(r, size, bold, italic)
    return p

def ch(text):           # chapter heading — centred, underline, 14pt
    p = _new_para(WD_ALIGN_PARAGRAPH.CENTER, 12, 6)
    r = p.add_run(text.upper())
    sf(r, 14, bold=True)
    r.font.underline = True

def sh(text):           # section heading — left, underline, 13pt
    p = _new_para(WD_ALIGN_PARAGRAPH.LEFT, 10, 4)
    r = p.add_run(text)
    sf(r, 13, bold=True)
    r.font.underline = True

def ssh(text):          # sub-section heading — left, bold, 12pt
    p = _new_para(WD_ALIGN_PARAGRAPH.LEFT, 8, 4)
    r = p.add_run(text)
    sf(r, 12, bold=True)

def bl(text):           # bullet
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    sf(r, 12)

def nb(text):           # numbered list
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    sf(r, 12)

def pb(): doc.add_page_break()

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '000000')
    pBdr.append(b); pPr.append(pBdr)

def table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        sf(r, 11, bold=True)
        c.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(txt))
            sf(r, 11)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def screenshot(fname, fig_label, caption_text, width_cm=14.5):
    """Insert a screenshot image with centred caption below it."""
    path = os.path.join(SHOT_DIR, fname)
    # image paragraph
    ip = _new_para(WD_ALIGN_PARAGRAPH.CENTER, 8, 2)
    r = ip.add_run()
    if os.path.exists(path):
        r.add_picture(path, width=Cm(width_cm))
    else:
        sf(r, 11, italic=True)
        r.text = f"[Image not found: {fname}]"
    # caption paragraph
    cp = _new_para(WD_ALIGN_PARAGRAPH.CENTER, 2, 12)
    cr = cp.add_run(f"Figure {fig_label}:  {caption_text}")
    sf(cr, 11, italic=True)

# ── Cover helper (bordered table cell) ───────────────────────────────────────
def cover_block():
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.paragraphs[0].clear()
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    return cell

def cp(cell, text, size=12, bold=False, sa=4):
    p = cell.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        r = p.add_run(text)
        sf(r, size, bold)
    return p

def cell_hline(cell):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'000000')
    pBdr.append(b)
    p._element.get_or_add_pPr().append(pBdr)

def remove_table_borders(tbl):
    for row in tbl.rows:
        for c in row.cells:
            tcPr = c._tc.get_or_add_tcPr()
            tcB  = OxmlElement('w:tcBorders')
            for bn in ['top','left','bottom','right','insideH','insideV']:
                b = OxmlElement(f'w:{bn}')
                b.set(qn('w:val'),'none')
                tcB.append(b)
            tcPr.append(tcB)

info_data = [
    ("Submitted by:",       "Om Prakash"),
    ("Roll No.:",           "[Roll Number]"),
    ("Branch:",             "Computer Science & Engineering"),
    ("Year / Semester:",    "3rd Year / 6th Semester"),
    ("Training Organisation:", "Novagito Technologies Pvt. Ltd."),
    ("Duration:",           "January 2026 – June 2026"),
]

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ═══════════════════════════════════════════════════════════════════════════════
cell = cover_block()
cp(cell, "INSTITUTE OF TECHNOLOGY & MANAGEMENT", 12, bold=True)
cp(cell, "(Affiliated to Dr. A.P.J. Abdul Kalam Technical University)", 11)
cell_hline(cell)
cp(cell, "Department of Computer Science & Engineering", 12, bold=True, sa=10)
cp(cell, "Internship / Training Project Report", 13, bold=True)
cp(cell, "Submitted in partial fulfilment of the requirements for the degree of", 11)
cp(cell, "Bachelor of Technology (B.Tech.)", 13, bold=True, sa=8)
cell_hline(cell)
cp(cell, "SCHOLARLY LIBRARY MANAGEMENT SYSTEM", 16, bold=True)
cp(cell, "A Full-Stack Web Application for Academic Library Operations", 12, sa=10)
cell_hline(cell)
it = cell.add_table(rows=6, cols=2)
it.style = 'Table Grid'
for i,(lbl,val) in enumerate(info_data):
    for j,txt in enumerate([lbl,val]):
        c = it.rows[i].cells[j]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        sf(r, 12, bold=(j==0))
remove_table_borders(it)
cell.add_paragraph()
cell_hline(cell)
cp(cell, "Academic Year: 2025–2026", 12, bold=True)
cp(cell, "June 2026", 12)

pb()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — INNER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
cell2 = cover_block()
cp(cell2, "INSTITUTE OF TECHNOLOGY & MANAGEMENT", 12, bold=True)
cp(cell2, "(Affiliated to Dr. A.P.J. Abdul Kalam Technical University)", 11)
cell_hline(cell2)
cp(cell2, "Department of Computer Science & Engineering", 12, bold=True, sa=10)
cp(cell2, "Internship / Training Project Report", 13, bold=True)
cp(cell2, "Submitted in partial fulfilment of the requirements for the degree of", 11)
cp(cell2, "Bachelor of Technology (B.Tech.)", 13, bold=True, sa=8)
cell_hline(cell2)
cp(cell2, "SCHOLARLY LIBRARY MANAGEMENT SYSTEM", 16, bold=True)
cp(cell2, "A Full-Stack Web Application for Academic Library Operations", 12, sa=10)
cell_hline(cell2)
it2 = cell2.add_table(rows=6, cols=2)
it2.style = 'Table Grid'
for i,(lbl,val) in enumerate(info_data):
    for j,txt in enumerate([lbl,val]):
        c = it2.rows[i].cells[j]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        sf(r, 12, bold=(j==0))
remove_table_borders(it2)
cell2.add_paragraph()
cell_hline(cell2)
cp(cell2, "Academic Year: 2025–2026", 12, bold=True)
cp(cell2, "June 2026", 12)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════
ch("Declaration")
mixed([("I, ",False,False,12),("Om Prakash",True,False,12),
       (", Roll No. ",False,False,12),("[Roll Number]",True,False,12),
       (", student of ",False,False,12),
       ("B.Tech. (Computer Science & Engineering), 6th Semester",True,False,12),
       (", at ",False,False,12),
       ("Institute of Technology & Management",True,False,12),
       (", hereby declare that the project report entitled ",False,False,12),
       ('"Scholarly Library Management System – A Full-Stack Web Application for Academic Library Operations"',False,True,12),
       (" has been prepared and submitted in partial fulfilment of the requirement of Semester Internship / Training, as prescribed by ",False,False,12),
       ("Dr. A.P.J. Abdul Kalam Technical University.",True,False,12)])
para("I further declare that:")
nb("The work contained in this report is original and has been done by me under the general supervision of my training supervisor at Novagito Technologies Pvt. Ltd. and the academic supervisor at my institute.")
nb("The work has not been submitted to any other university or institution for the award of any degree, diploma, or certificate.")
nb("I have followed the guidelines provided by the institute (Annexure VI) while writing this report.")
nb("Wherever I have used materials (data, theoretical analysis, figures, or text) from other sources, I have given due credit and full reference details in the bibliography.")
doc.add_paragraph()
sig = doc.add_table(rows=1, cols=2)
sig.style = 'Table Grid'
remove_table_borders(sig)
lc, rc = sig.rows[0].cells[0], sig.rows[0].cells[1]
lc.paragraphs[0].clear(); r=lc.paragraphs[0].add_run("Place: [City Name]"); sf(r,12)
lp=lc.add_paragraph(); r=lp.add_run("Date: June 2026"); sf(r,12)
rc.paragraphs[0].clear()
for _ in range(3): rc.add_paragraph()
rp = rc.add_paragraph()
rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pBdr = OxmlElement('w:pBdr')
t = OxmlElement('w:top')
t.set(qn('w:val'),'single'); t.set(qn('w:sz'),'6')
t.set(qn('w:space'),'1'); t.set(qn('w:color'),'000000')
pBdr.append(t); rp._element.get_or_add_pPr().append(pBdr)
r=rp.add_run("Om Prakash"); sf(r,12,bold=True)
rp2=rc.add_paragraph(); rp2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=rp2.add_run("Roll No.: [Roll Number]"); sf(r,11)
rp3=rc.add_paragraph(); rp3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=rp3.add_run("B.Tech. CSE, 6th Semester"); sf(r,11)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════════════════════
ch("Acknowledgement")
para("I express my heartfelt gratitude to all those who guided, supported, and encouraged me throughout the course of this internship and in the preparation of this project report.")
mixed([("First and foremost, I am deeply thankful to ",False,False,12),
       ("Novagito Technologies Pvt. Ltd.",True,False,12),
       (" for providing me with the opportunity to undertake my semester internship and to work on a meaningful, real-world full-stack web development project. The experience gained at the organisation has been invaluable.",False,False,12)])
mixed([("I extend my sincere gratitude to my ",False,False,12),
       ("Industry Supervisor",True,False,12),
       (" at Novagito Technologies for expert guidance, constructive feedback, and constant encouragement at every stage of project development.",False,False,12)])
mixed([("I am equally grateful to my ",False,False,12),
       ("Faculty Supervisor",True,False,12),
       (" and all faculty members of the ",False,False,12),
       ("Department of Computer Science & Engineering",True,False,12),
       (", ",False,False,12),
       ("Institute of Technology & Management",True,False,12),
       (", for academic guidance and for facilitating this internship programme.",False,False,12)])
para("I also thank the Head of Department and the Principal for their continuous support and for providing the required academic infrastructure.")
para("Special thanks to my fellow interns and colleagues at Novagito Technologies whose discussions and code reviews significantly improved the quality of the final application.")
para("Last but not least, I am deeply indebted to my family for their patience, encouragement, and moral support.")
doc.add_paragraph()
ap = doc.add_paragraph()
ap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for t in ["Om Prakash","Roll No.: [Roll Number]","B.Tech. CSE, 6th Semester","Institute of Technology & Management"]:
    r = ap.add_run(t+"\n"); sf(r,12,bold=(t=="Om Prakash"))
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
ch("Contents")
toc_items=[
    ("1.","Introduction","6",False),
    ("1.1","Background","6",True),
    ("1.2","Problem Statement","7",True),
    ("1.3","Scope of the Project","7",True),
    ("2.","Objective","8",False),
    ("3.","About the Organisation","9",False),
    ("4.","Project / Work Description","10",False),
    ("4.1","Technology Stack","10",True),
    ("4.2","System Architecture","11",True),
    ("4.3","Database Design","13",True),
    ("4.4","API Design","15",True),
    ("4.5","Frontend Architecture","17",True),
    ("4.6","User Interface Design System","18",True),
    ("5.","Achievement of Work","19",False),
    ("5.1","Authentication and User Registration","19",True),
    ("5.2","Material Catalogue","20",True),
    ("5.3","User Dashboard","21",True),
    ("5.4","Activity Feed and Notifications","22",True),
    ("5.5","User Profile and Saved Materials","22",True),
    ("5.6","Admin Analytics Dashboard","23",True),
    ("5.7","Borrow Request Management","24",True),
    ("5.8","User and Material Management","25",True),
    ("5.9","Security Implementation","26",True),
    ("5.10","Role-Based Access Control","27",True),
    ("5.11","Technical Challenges and Solutions","27",True),
    ("6.","Conclusions and Future Scope","28",False),
    ("7.","Impediments, Reflections and Suggestions","29",False),
    ("8.","References","31",False),
]
for num,title,pg,is_sub in toc_items:
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
    p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
    if is_sub: p.paragraph_format.left_indent=Cm(1)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.5),leader=1)
    r1=p.add_run(f"{num}  {title}"); r2=p.add_run("\t"); r3=p.add_run(pg)
    sf(r1,12,bold=(not is_sub)); sf(r2,12); sf(r3,12)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
ch("1. Introduction")
sh("1.1 Background")
para("Libraries are the cornerstone of academic institutions, serving as repositories of knowledge and facilitating learning for students and researchers alike. Traditionally, library operations — including cataloguing, borrowing, returning, and tracking materials — have been managed through manual or semi-automated systems. These approaches are prone to inefficiencies such as record duplication, delayed processing of borrow requests, lack of real-time availability information, and limited accessibility beyond physical premises.")
para("With the rapid advancement of web technologies and the widespread adoption of the internet across academic campuses, there is a growing demand for digital library management systems that can streamline operations, improve user experience, and provide instant access to library information from any device.")
mixed([("The ",False,False,12),("Scholarly Library Management System",True,False,12),
       (" was conceived and developed as part of this internship project to address precisely these needs. It is a full-stack web application built on React, Node.js, Express, and MongoDB, designed to replace traditional library workflows with an efficient, scalable, and user-friendly digital solution.",False,False,12)])

sh("1.2 Problem Statement")
para("Academic institutions face several common challenges in library management:")
bl("Manual borrow/return tracking — Registers and spreadsheets are error-prone, difficult to audit, and inaccessible remotely.")
bl("Limited search capability — Students cannot quickly locate materials by title, author, or subject without physically browsing shelves.")
bl("No real-time availability — Students have no way to know if a book is available before visiting the library.")
bl("Inefficient communication — Librarians have no efficient channel to notify students about approvals, rejections, or due dates.")
bl("No digital access — Notes and PYQ documents are not accessible digitally, requiring physical presence.")
bl("Poor data for decision-making — Librarians lack analytics on material popularity and user engagement.")

sh("1.3 Scope of the Project")
bl("A secure, JWT-based authentication system with email verification and role management.")
bl("A material catalogue supporting books, notes, and past-year question papers with full-text search.")
bl("A borrow workflow covering request submission, librarian approval/rejection, renewals, and return tracking.")
bl("An in-browser document viewer for digital materials.")
bl("A librarian-facing administrative dashboard with analytics, user management, and material management.")
bl("A notification/activity feed for users to track their library interactions.")
bl("A responsive dark/light themed web interface following Material Design 3 principles.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — OBJECTIVE
# ═══════════════════════════════════════════════════════════════════════════════
ch("2. Objective")
sh("2.1 Primary Objectives")
nb("To design and develop a full-stack web application that digitises and streamlines all key operations of an academic library.")
nb("To implement a secure, role-based authentication system using JWT and bcrypt password hashing, supporting patron, librarian, and super-administrator roles.")
nb("To create a comprehensive material catalogue covering books, academic notes, and past-year question papers with full-text search and real-time availability tracking.")
nb("To implement an end-to-end borrow management workflow — borrow requests, approval/rejection, renewals, and return processing — with automated notifications.")
nb("To develop a librarian-facing administrative interface with analytics dashboards, user management, material management, and borrow history audit trails.")
nb("To build a user-friendly patron interface for browsing, saving, borrowing, and reading digital materials in-browser.")
nb("To implement an email notification system using Nodemailer for account verification (OTP and token-based) and borrow-related alerts.")
sh("2.2 Learning Objectives")
nb("To gain hands-on experience in full-stack JavaScript development using the MERN-style stack.")
nb("To learn and apply RESTful API design principles in building a production-grade backend service.")
nb("To understand state management patterns in React — Context API, custom hooks, and session-based authentication.")
nb("To apply Google's Material Design 3 system in building an accessible, visually consistent interface.")
nb("To practise professional software development workflows including modular code organisation and role-based access control.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — ABOUT THE ORGANISATION
# ═══════════════════════════════════════════════════════════════════════════════
ch("3. About the Organisation")
sh("3.1 Novagito Technologies Pvt. Ltd.")
para("Novagito Technologies Pvt. Ltd. is a software development company specialising in full-stack web application development, enterprise software solutions, and digital product engineering. The company builds scalable, user-centric software products for clients across education, e-commerce, and enterprise sectors.")
para("Novagito Technologies works with modern web technologies including the JavaScript/Node.js ecosystem, React-based frontends, cloud infrastructure (AWS, GCP), and NoSQL databases such as MongoDB. The company follows agile development methodologies emphasising iterative delivery, code quality, and continuous integration practices.")
sh("3.2 Internship Environment")
para("The internship was conducted under the Software Development team at Novagito Technologies. The intern was assigned to build a modern Library Management System as a portfolio-grade demonstration of full-stack development capabilities using the latest stable versions of React, Node.js, and MongoDB.")
bl("Weekly sprint planning — features broken into tasks with defined priorities.")
bl("Daily stand-up reviews — progress, blockers, and next steps discussed.")
bl("Code review sessions — senior developers provided feedback on structure, security, and UI/UX.")
bl("Documentation requirements — all API endpoints and data models documented as part of development.")
sh("3.3 Tools and Infrastructure")
table(["Category","Tool / Service","Purpose"],
    [("Version Control","Git / GitHub","Source code management"),
     ("IDE","Visual Studio Code","Code editing and debugging"),
     ("Database","MongoDB Atlas","Cloud-hosted MongoDB cluster"),
     ("API Testing","Postman","Backend API testing"),
     ("Email Testing","Mailtrap / Gmail SMTP","Nodemailer email tests"),
     ("Communication","Slack / Zoom","Team communication"),
     ("Project Management","Notion","Task tracking and sprints")],
    col_widths=[4,4,7])
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — PROJECT / WORK DESCRIPTION
# ═══════════════════════════════════════════════════════════════════════════════
ch("4. Project / Work Description")
para("The Scholarly Library Management System is a full-stack web application built in a monorepo structure — frontend/ (React + Vite) and backend/ (Node.js + Express). This chapter covers technology choices, architecture, database design, API structure, and the UI design system.")

sh("4.1 Technology Stack")
ssh("4.1.1 Frontend")
table(["Technology","Version","Role"],
    [("React","19.2.6","UI component library and SPA framework"),
     ("Vite","8.0.12","Build tool and development server"),
     ("React Router DOM","7.15.0","Client-side routing"),
     ("TailwindCSS","3.4.19","Utility-first CSS framework"),
     ("TypeScript","~6.0.2","Static type checking"),
     ("Google Material Symbols","—","Icon font (Outlined variant)")],
    col_widths=[4,4,7])

ssh("4.1.2 Backend")
table(["Technology","Version","Role"],
    [("Node.js","LTS","JavaScript runtime"),
     ("Express.js","4.19.2","HTTP server and REST API"),
     ("Mongoose","8.4.0","MongoDB ODM"),
     ("MongoDB","Atlas","NoSQL cloud database"),
     ("jsonwebtoken","9.0.2","Stateless JWT authentication"),
     ("bcryptjs","2.4.3","Password hashing (10 salt rounds)"),
     ("Multer","2.1.1","File upload handling"),
     ("Nodemailer","8.0.10","Transactional email (OTP, alerts)"),
     ("express-validator","7.1.0","Request body validation"),
     ("CORS","2.8.5","Cross-origin request control")],
    col_widths=[4,4,7])

sh("4.2 System Architecture")
para("The system follows a three-tier client-server architecture:")
bl("Client Tier — React SPA (Vite) running in the user's browser, communicating via HTTP REST API.")
bl("Application Tier — Node.js + Express server (port 5000) exposing /api endpoints with JWT middleware and business logic controllers.")
bl("Data Tier — MongoDB Atlas providing cloud-hosted NoSQL storage with Mongoose ODM for schema management and full-text indexing.")

ssh("4.2.1 Context Provider Hierarchy (Frontend)")
para("The React app wraps all components in nested providers: AuthContext → ThemeContext → ProfileContext → ActivityContext → UserLibraryContext → RouterProvider. Authentication state is universally available and theme is applied at root level before any rendering occurs.")

ssh("4.2.2 Request–Response Cycle")
para("(1) React component calls services/api.js → (2) JWT attached as Authorization: Bearer header → (3) Express protect() middleware validates JWT → (4) Controller executes business logic → (5) MongoDB read/write via Mongoose → (6) JSON response returned → (7) React updates state, ActionToast shows feedback.")

sh("4.3 Database Design")
ssh("4.3.1 User Collection")
table(["Field","Type","Description"],
    [("name","String (required)","Full name"),
     ("email","String (unique)","Login identifier"),
     ("password","String (hashed)","Bcrypt 10-round hash"),
     ("role","user | admin | superadmin","Access level"),
     ("picture, phone, dept, course, year, bio","String","Extended profile"),
     ("suspended, isDeactivated","Boolean","Account status flags"),
     ("otpCode, otpExpiry","String / Date","Email OTP verification"),
     ("isEmailVerified, emailVerifyToken","Boolean / String","Email verification state")],
    col_widths=[4.5,4.5,6])
ssh("4.3.2 Material Collection")
table(["Field","Type","Description"],
    [("title","String (required)","Material title"),
     ("author","String","Author or creator"),
     ("kind","book | notes | pyq","Material type"),
     ("categoryLabel","String","Subject / category"),
     ("coverUrl, fileUrl","String","Cover image and document paths"),
     ("rating","Number (0–5)","Average rating"),
     ("totalCopies, availableCopies","Number","Copy inventory"),
     ("addedBy","ObjectId → User","Librarian who added it"),
     ("Text Index","—","On title, author, categoryLabel")],
    col_widths=[4.5,4.5,6])
ssh("4.3.3 BorrowRequest Collection")
table(["Field","Type","Description"],
    [("user","ObjectId → User","Patron requesting the borrow"),
     ("material","ObjectId → Material","Material being borrowed"),
     ("status","pending|approved|rejected|returned|cancelled","Borrow lifecycle state"),
     ("requestedAt, approvedAt, returnedAt, dueDate","Date","Key lifecycle timestamps"),
     ("renewCount","Number","Number of renewals"),
     ("processedBy","ObjectId → User","Admin who processed it")],
    col_widths=[4.5,5,5.5])
para("SavedItem stores user bookmarks with a unique compound index on (user, material). Activity stores user notification events with kind, title, detail, icon styling, and dismissed flag.")

sh("4.4 API Design")
table(["Method","Endpoint","Description","Auth"],
    [("POST","/api/auth/signup","Register new user","Public"),
     ("POST","/api/auth/login","Authenticate and receive JWT","Public"),
     ("POST","/api/auth/verify-email","Verify email via OTP or token","Public"),
     ("GET","/api/materials","List materials (search + filter)","Public"),
     ("POST","/api/materials","Create material + file upload","Admin"),
     ("PATCH","/api/materials/:id","Update material","Admin"),
     ("DELETE","/api/materials/:id","Delete material","Admin"),
     ("POST","/api/borrow/request","Submit borrow request","User"),
     ("GET","/api/borrow/all","All borrows with filters","Admin"),
     ("PATCH","/api/borrow/:id","Approve/reject/renew/return","Admin"),
     ("GET","/api/users","List all users","Admin"),
     ("PATCH","/api/users/:id/suspend","Suspend user account","Admin"),
     ("GET","/api/activity","User notification feed","User"),
     ("POST","/api/saved","Bookmark a material","User"),
     ("POST","/api/upload","Upload file (profile/document)","User/Admin")],
    col_widths=[2,4.5,6.5,2])

sh("4.5 Frontend Architecture")
para("Source is organised into: components/ (layout, navigation, ui), contexts/ (Auth, Theme, Profile, Activity, UserLibrary), hooks/ (useMaterials, useUsers), pages/ (auth/, user/, admin/), routes/ (appPaths.js), services/ (api.js), utils/ (timeAgo.js).")
para("React Router DOM v7 divides routes into: public (/login, /signup, /verify-email), patron (/user/*) protected by role 'user', and admin (/admin/*) protected by role 'admin'/'superadmin', wrapped in AdminLayout with sidebar navigation.")

sh("4.6 User Interface Design System")
table(["Token","Light Mode","Dark Mode","Usage"],
    [("--md-primary","#1a1a2e","#f0ede5 (warm white)","Headings, primary text"),
     ("--md-secondary","#8b6f3e","#e9c176 (gold)","Accents, interactive elements"),
     ("--md-surface","#f7f9fb","#0c0c14 (near black)","Page background"),
     ("--md-surface-container","#eef0f3","#161624","Card backgrounds"),
     ("--md-on-surface-variant","#6b5d3e","#c0b090 (gold-tinted)","Secondary body text"),
     ("--md-error","#ba1a1a","#ffb4ab","Error states")],
    col_widths=[4,3.5,3.5,4])
para("Display/headline typeface: Manrope (700–800 weight). Body/UI typeface: Work Sans (400–600). Icons: Google Material Symbols Outlined. Dark mode applied via .dark class on <html> root — all CSS tokens switch automatically, zero per-component dark: prefixes needed.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — ACHIEVEMENT OF WORK (with inline screenshots)
# ═══════════════════════════════════════════════════════════════════════════════
ch("5. Achievement of Work")
para("This chapter details all modules and features developed during the six-month internship, with screenshots of the live application embedded alongside each feature's description. The Scholarly Library Management System was built iteratively with features progressively added and refined based on supervisor feedback and testing.")

# ── 5.1 Authentication ────────────────────────────────────────────────────────
sh("5.1 Authentication and User Registration")
ssh("5.1.1 Login Page")
para("The Login Page (/login) allows registered students and librarians to authenticate with their institutional email and password. The page features the application brand on the left with a hero statement, and a clean credential form on the right with inline validation and a 'Forgot Password' link. JWT token is stored in sessionStorage on successful login.")
screenshot("01_login_page.png","5.1","Login Page — Dark-themed authentication screen with brand hero and credential form.")

pb()
ssh("5.1.2 Signup and Email Verification")
para("The Signup Page (/signup) collects name, email, and password. After form submission, the backend generates a 6-digit OTP and dispatches a styled HTML verification email via Nodemailer. The student must verify their email before accessing library features. Both OTP entry and a verification-link pathway are supported, with OTP expiry enforced for security.")
screenshot("02_signup_page.png","5.2","Signup / Registration Page — New user account creation with email verification flow.")

pb()
# ── 5.2 Material Catalogue ────────────────────────────────────────────────────
sh("5.2 Material Catalogue")
para("The Catalogue Page (/user/catalog) presents all library materials — books, notes, and past-year question papers (PYQ) — in a responsive card grid. Each card displays the cover image, title, author, colour-coded kind badge, star rating, and availability status. A real-time debounced search queries MongoDB's text index on title, author, and categoryLabel fields. Tabs filter results by material type (All / Books / Notes / PYQ).")
para("From any material card, students navigate to the Material Detail Page which displays full description, availability, and provides options to submit a borrow request, join a waitlist, save the material, view similar materials, or open the in-browser document viewer.")
screenshot("04_material_catalog.png","5.3","Material Catalogue — Card grid showing books, notes, and PYQ materials with search and filter controls.")

pb()
# ── 5.3 User Dashboard ────────────────────────────────────────────────────────
sh("5.3 User Dashboard")
para("The User Dashboard (/user/dashboard) is the patron's home page. It provides an at-a-glance view of the student's current borrow status (approved borrows with due dates), featured materials for discovery, and quick-access navigation to the catalogue, activity feed, and profile. The dashboard is personalised based on the logged-in user's borrow and save history, encouraging continued engagement with the library.")
screenshot("03_user_dashboard.png","5.4","User Dashboard — Patron home page showing active borrow status and featured materials.")

pb()
# ── 5.4 Activity Feed ────────────────────────────────────────────────────────
sh("5.4 Activity Feed and Notifications")
para("The Activity Page (/user/activity) presents a chronological notification feed covering all library events for the user — borrow approvals, rejections, returns, renewals, due-date alerts, and library announcements. Each entry displays a colour-coded icon (by event kind), a relative timestamp computed by the timeAgo.js utility (e.g. '2 hours ago'), the event title and detail text, and a dismiss button. Unread notification counts are reflected in the RouteNav badge.")
screenshot("05_activity_feed.png","5.5","Activity Feed — Chronological notification log showing all library events for the user.")

pb()
# ── 5.5 Profile and Saved Materials ──────────────────────────────────────────
sh("5.5 User Profile and Saved Materials")
ssh("5.5.1 User Profile")
para("The Profile Page (/user/profile) allows users to update personal information — name, phone number, department, course year, and bio — and upload a profile picture via the /api/upload endpoint. Profile data is managed by ProfileContext and reflected in real-time across the navigation bar and all user-facing components.")
screenshot("06_user_profile.png","5.6","User Profile Page — Personal information editing with profile picture upload.")

pb()
ssh("5.5.2 Saved Materials")
para("The Save History Page (/user/save-history) lists all materials bookmarked by the user with their cover, title, kind, and category. Each bookmark can be navigated to for full detail or removed with a single action. Saved items are persisted in the SavedItem MongoDB collection with a unique compound index on (user, material) preventing duplicate saves.")
screenshot("07_saved_materials.png","5.7","Saved / Bookmarked Materials — User's personal shelf of saved library items.")

pb()
# ── 5.6 Admin Analytics ──────────────────────────────────────────────────────
sh("5.6 Admin Analytics Dashboard")
para("The Admin Analytics Page (/admin/analytics) provides the librarian with a real-time statistical overview of the entire library system. Key metrics displayed include: total registered users, total materials in the catalogue, active (approved) borrows, pending borrow requests awaiting action, total save events, and returned materials count. These metrics are fetched live from MongoDB aggregations, enabling data-driven decisions about collection development and staffing.")
screenshot("08_admin_analytics.png","5.8","Admin Analytics Dashboard — Real-time library statistics showing users (5), saved items (22), and borrow metrics.")

pb()
# ── 5.7 Borrow Management ────────────────────────────────────────────────────
sh("5.7 Borrow Request Management")
para("The Admin Borrows Page (/admin/borrows) is the operational centre of the librarian interface. It lists all borrow requests filterable by status (pending, approved, rejected, returned, cancelled). For each request the librarian sees: patron name, material title, request date, and current status. Actions available per request:")
bl("Approve — sets status to 'approved', records approvedAt and dueDate, decrements material's availableCopies, and creates an Activity notification for the student.")
bl("Reject — with optional reason, notifies the student via their activity feed.")
bl("Mark as Returned — sets status to 'returned', records returnedAt, and increments availableCopies.")
bl("Process Renewal — extends the dueDate for an already-approved borrow.")
screenshot("09_admin_borrows.png","5.9","Borrow Management — Librarian interface for approving, rejecting, and processing all borrow requests.")

pb()
para("The Admin Borrow History Page (/admin/borrow-history) provides a complete, filterable audit trail of all borrow transactions across all users, including historical records of returned and rejected requests.")
screenshot("13_borrow_history.png","5.10","Borrow History — Complete audit trail of all past borrow transactions across all users.")

pb()
# ── 5.8 User and Material Management ─────────────────────────────────────────
sh("5.8 User and Material Management")
ssh("5.8.1 User Management")
para("The Admin Users Page (/admin/users) provides a paginated, searchable list of all registered users displaying role, email verification status, and account flags. Librarians can: create new user accounts directly via AdminUserFormPage (bypassing self-registration); view and edit user profiles with role assignment; and suspend or unsuspend accounts without deleting the audit trail.")
screenshot("10_admin_users.png","5.11","User Management — Admin view of all registered users with role, status, and action controls.")

pb()
ssh("5.8.2 Material Management")
para("The Materials Page (/admin/materials) lists all library items with inline edit and delete controls. The Add Material Page (/admin/materials/add) provides a comprehensive form for adding new books, notes, or PYQs including title, author, category, kind, description, copy count, cover image, and document file upload via Multer.")
screenshot("11_admin_materials.png","5.12","Materials Management — Admin catalogue view with edit, delete, and availability controls.")

pb()
ssh("5.8.3 Add Material Form")
para("The Add Material form demonstrates the complete material creation workflow, handling both metadata (title, author, category, kind, description) and file uploads (cover image, document) in a single multipart/form-data request. The Update Material page provides the same form pre-populated with existing data for inline editing.")
screenshot("12_add_material_form.png","5.13","Add New Material Form — Librarian form for adding books, notes, or PYQs with file upload.")

pb()
# ── 5.9 Security ─────────────────────────────────────────────────────────────
sh("5.9 Security Implementation")
bl("Password Hashing: All passwords hashed with bcryptjs at 10 salt rounds. Plain-text passwords never stored in the database.")
bl("JWT Authentication: Signed JWT issued on login, stored in sessionStorage (cleared on tab close, reducing XSS exposure). Auto-injected on every API request via services/api.js.")
bl("Role-Based Middleware: Backend requireRole() enforces admin/superadmin access on all administrative endpoints. Frontend ProtectedRoute enforces the same at UI routing level.")
bl("Email Verification: New accounts must verify email before accessing library features, preventing fake registrations.")
bl("Input Validation: All API request bodies validated via express-validator to prevent injection attacks and malformed data.")
bl("CORS Restriction: Express only accepts cross-origin requests from whitelisted origins.")
bl("File Upload Limits: 5 MB body limit on Express JSON parser; Multer enforces MIME type and size restrictions.")
bl("Account Suspension: Admins can suspend accounts without deletion, preserving audit trails.")

sh("5.10 Role-Based Access Control")
table(["Feature / Action","User (Patron)","Admin (Librarian)","Superadmin"],
    [("Browse/Search Materials","Yes","Yes","Yes"),
     ("Submit Borrow Request","Yes","No","No"),
     ("Save/Bookmark Material","Yes","No","No"),
     ("View Own Activity Feed","Yes","Yes","Yes"),
     ("Approve / Reject Borrows","No","Yes","Yes"),
     ("Add / Edit / Delete Material","No","Yes","Yes"),
     ("Manage Users","No","Yes","Yes"),
     ("View Analytics Dashboard","No","Yes","Yes"),
     ("Suspend User Accounts","No","Yes","Yes"),
     ("Assign Admin Role","No","No","Yes")],
    col_widths=[6,3,4,3])

sh("5.11 Technical Challenges and Solutions")
ssh("Challenge 1 — Theme Switching Without Flash of Unstyled Content (FOUC)")
para("Solution: A synchronous inline script in index.html reads localStorage before the first React render and applies the .dark class to the HTML root. TailwindCSS's darkMode: 'class' combined with CSS custom properties in index.css ensures all components respond automatically.")
ssh("Challenge 2 — Auth State Synchronisation Across Components")
para("Solution: AuthContext implemented using React's useSyncExternalStore hook, subscribing to sessionStorage events so logout is reflected instantly across all mounted components without manual state propagation.")
ssh("Challenge 3 — MongoDB Text Search with Category Filtering")
para("Solution: The materials list endpoint dynamically builds the query object — applying $text search only when a search query is present and appending kind filters only when specified, satisfying MongoDB's constraint that $text cannot be combined with other indexed field queries in all cases.")
ssh("Challenge 4 — File Upload with Metadata in a Single Request")
para("Solution: Multer configured for multipart/form-data on the backend; frontend uses FormData API to bundle file and metadata in one POST. Error handling covers size limits, invalid MIME types, and partial upload failures.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — CONCLUSIONS & FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════════════════════
ch("6. Conclusions and Future Scope of Work")
sh("6.1 Conclusions")
para("The Scholarly Library Management System was successfully designed, developed, and tested over six months (January–June 2026) at Novagito Technologies Pvt. Ltd. The application achieves all stated objectives and demonstrates a production-grade full-stack JavaScript implementation.")
nb("The MERN-style stack proved highly effective — JavaScript uniformity across frontend and backend reduced context-switching and enabled efficient code reuse.")
nb("React 19 with Context API provides lightweight, maintainable state management for this scale without requiring Redux or Zustand.")
nb("Material Design 3 as a design system significantly reduced design decision fatigue and produced a professional-quality dark-mode UI.")
nb("JWT-based RBAC is robust and scalable for multi-role applications when tokens are stored securely and expiry is enforced.")
nb("MongoDB's document model and text indexing enabled flexible schemas and full-text search without a dedicated search engine.")
nb("The project demonstrated that a student developer can produce a commercially-quality full-stack application within an internship timeframe given an iterative, feedback-driven development process.")

sh("6.2 Future Scope")
ssh("Short-Term (0–6 months)")
bl("Mobile responsive navigation — hamburger menu and bottom nav bar for mobile screens.")
bl("Server-side pagination and infinite scroll for large material collections.")
bl("Automated overdue fine calculation and email reminders X days before due date.")
bl("Advanced analytics with Recharts/Chart.js visualisations of borrow trends.")
ssh("Medium-Term (6–18 months)")
bl("Progressive Web App (PWA) for offline browsing and push notifications.")
bl("Barcode/QR code scanning for quick material lookup and borrow processing.")
bl("Collaborative filtering recommendation engine based on borrow and save history.")
bl("Google Books API integration to auto-populate material metadata by ISBN.")
ssh("Long-Term (18+ months)")
bl("Digital Rights Management (DRM) for time-limited ebook lending.")
bl("AI-powered study assistant (LLM integration) for Q&A over digital documents.")
bl("RFID hardware integration for self-service borrow/return stations.")
bl("Cloud object store migration (AWS S3 / Cloudflare R2) for scalable file delivery.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 — IMPEDIMENTS / REFLECTIONS / SUGGESTIONS
# ═══════════════════════════════════════════════════════════════════════════════
ch("7. Impediments / Difficulties, Reflections and Suggestions")
sh("7.1 Impediments and Difficulties Faced")
ssh("Technical Difficulties")
para("a) React 19 and useSyncExternalStore: Adapting to React 19's concurrent features and implementing the useSyncExternalStore hook for sessionStorage synchronisation in AuthContext without causing infinite re-renders required dedicated study and iterative debugging.")
para("b) TailwindCSS + MD3 Token Mapping: Configuring Google's Material Design 3 colour tokens as CSS custom properties that interact correctly with TailwindCSS's JIT compiler and the .dark class cascade required multiple configuration iterations.")
para("c) JWT Expiry Handling: Implementing a global 401 interceptor in api.js that triggers logout without cascading re-render loops required careful state management design.")
para("d) Mongoose Population: Chaining populate() calls with field projections to exclude sensitive data (e.g., password hashes) from API responses required close security review.")
para("e) File Upload Architecture: Deciding between local storage (backend/uploads/) and cloud storage involved trade-offs between deployment simplicity and production scalability. Local storage was chosen for internship scope with cloud migration flagged as a future enhancement.")
ssh("Non-Technical Difficulties")
para("a) Scope Management: The open-ended brief led to feature creep. Weekly sprint planning introduced MVP-first discipline, ensuring core features were completed before enhancements were considered.")
para("b) Balancing Obligations: Managing concurrent academic responsibilities alongside internship deliverables required a Notion task board and dedicated time blocks for each.")

sh("7.2 Reflection on Learning and Value Addition")
ssh("Technical Skills")
bl("Full-Stack JavaScript Development: End-to-end hands-on experience across MongoDB, Express, React, and TailwindCSS — difficult to replicate in academic settings.")
bl("RESTful API Design: Resource-oriented endpoint design with correct HTTP semantics (GET/POST/PATCH/DELETE) and status codes.")
bl("Authentication and Security: JWT, bcrypt, RBAC, email verification, CORS, and input validation in a real application.")
bl("React Patterns: Context API, useSyncExternalStore, custom hooks (useMaterials, useUsers), and React Router v7.")
bl("Design Systems: Token-driven MD3 theming — scalable UI engineering beyond ad-hoc styling.")
ssh("Professional Skills")
bl("Agile / Sprint Development: Weekly sprints, stand-ups, and retrospectives — industry Agile practice first-hand.")
bl("Code Review Culture: Receiving and incorporating senior developer feedback accelerated growth and developed a discipline for readable, maintainable code.")
bl("Technical Documentation: Writing API documentation, sprint updates, and this report improved technical communication skills essential in professional teams.")
bl("Problem-Solving Under Deadlines: Pragmatic solutions under time constraints — a skill that cannot be self-taught.")

sh("7.3 Suggestions")
ssh("Suggestions for the Project")
bl("Migrate file storage to AWS S3 or Cloudflare R2 before production deployment — local storage is not durable in containerised environments.")
bl("Add Jest (backend) and React Testing Library (frontend) test suites to improve codebase maintainability.")
bl("Introduce API versioning (/api/v1/*) from the outset to allow future breaking changes without disrupting consumers.")
bl("Add express-rate-limit on authentication endpoints to prevent brute-force attacks.")
bl("Implement a CDN for cover images and document files for faster global delivery.")
ssh("Suggestions for the Internship Programme")
bl("Provide a starter kit/boilerplate so interns focus on features rather than spending weeks on scaffolding.")
bl("Organise bi-weekly structured code review sessions where interns present work to senior engineers.")
bl("Introduce DevOps practices — CI/CD pipelines, Docker, basic cloud deployment — even briefly.")
bl("Provide a shared test MongoDB Atlas cluster from day one to avoid database provisioning delays.")
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
ch("8. References")
refs=[
    "Meta Platforms Inc., React Documentation — React 19 Release, react.dev. [Accessed: June 2026]",
    "Vite Team, Vite — Next Generation Frontend Tooling, vite.dev. [Accessed: June 2026]",
    "Remix Software, React Router DOM v7 Documentation, reactrouter.com. [Accessed: June 2026]",
    "Tailwind Labs, Tailwind CSS v3 Documentation, tailwindcss.com/docs. [Accessed: June 2026]",
    "OpenJS Foundation, Node.js Documentation, nodejs.org/en/docs. [Accessed: June 2026]",
    "StrongLoop / IBM, Express.js 4.x API Reference, expressjs.com. [Accessed: June 2026]",
    "MongoDB Inc., MongoDB Manual — CRUD Operations and Text Search, mongodb.com/docs. [Accessed: June 2026]",
    "Mongoose Ltd., Mongoose ODM v8 Documentation, mongoosejs.com/docs. [Accessed: June 2026]",
    "Auth0, JWT.io — JSON Web Tokens Introduction, jwt.io/introduction. [Accessed: June 2026]",
    "npm Inc., bcryptjs — npm Package, npmjs.com/package/bcryptjs. [Accessed: June 2026]",
    "Nodemailer Team, Nodemailer Documentation, nodemailer.com/about. [Accessed: June 2026]",
    "Google LLC, Material Design 3 — Design System, m3.material.io. [Accessed: June 2026]",
    "Google Fonts, Manrope and Work Sans Typefaces, fonts.google.com. [Accessed: June 2026]",
    "npm Inc., Multer — File Upload Middleware, npmjs.com/package/multer. [Accessed: June 2026]",
    "Alex Nicol, express-validator v7 Documentation, express-validator.github.io. [Accessed: June 2026]",
    "npm Inc., dotenv, npmjs.com/package/dotenv. [Accessed: June 2026]",
    "R. Elmasri, S.B. Navathe, Fundamentals of Database Systems, 7th Ed., Pearson, 2015.",
    "D. Flanagan, JavaScript: The Definitive Guide, 7th Ed., O'Reilly, 2020.",
    "A. Osmani, Learning JavaScript Design Patterns, 2nd Ed., O'Reilly, 2023.",
    "OWASP Foundation, OWASP Top Ten — Web Application Security Risks, owasp.org. [Accessed: June 2026]",
]
for i,ref in enumerate(refs,1):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent=Cm(0.8); p.paragraph_format.first_line_indent=Cm(-0.8)
    r=p.add_run(f"[{i}]  {ref}"); sf(r,12)

# ── Final line ─────────────────────────────────────────────────────────────────
doc.add_paragraph(); hline()
ep=doc.add_paragraph()
ep.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
ep.paragraph_format.space_before=Pt(6)
r=ep.add_run("— End of Project Report —"); sf(r,12,bold=True)
sp=doc.add_paragraph()
sp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Scholarly Library Management System  |  Om Prakash  |  B.Tech. CSE  |  June 2026")
sf(r,11,italic=True)

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
